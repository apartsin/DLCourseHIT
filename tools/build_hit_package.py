# -*- coding: utf-8 -*-
"""Generate the HIT catalogue package for Introduction to Deep Learning, mirroring the
HITCatalogueExamples templates: English syllabus, Hebrew syllabus, rationale,
catalogue summary, and committee questionnaire. Output: hit-catalogue/*.docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT = os.path.join(ROOT, "hit-catalogue")
HE_FONT, EN_FONT = "Arial", "Arial"          # match the on-disk HIT examples (Arial throughout)
LOGO = os.path.join(os.path.dirname(os.path.abspath(__file__)), "hit-logo.png")

def new_doc():
    """A4 document pre-loaded with the official HIT logo letterhead (matches the HIT examples)."""
    d = Document()
    sec = d.sections[0]
    sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)          # A4
    sec.left_margin = sec.right_margin = Cm(3.0)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    hp = sec.header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hp.add_run().add_picture(LOGO, width=Inches(1.7))
    return d

def _set_rtl(p):
    pPr = p._p.get_or_add_pPr()
    b = OxmlElement('w:bidi'); pPr.append(b)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

def _set_table_rtl(t):
    """Make a table right-to-left: the first logical column renders on the right."""
    tblPr = t._tbl.tblPr
    if tblPr.find(qn('w:bidiVisual')) is None:
        tblPr.append(OxmlElement('w:bidiVisual'))

def set_doc_rtl(doc):
    """Make the whole document section right-to-left (Hebrew docs)."""
    sectPr = doc.sections[0]._sectPr
    if sectPr.find(qn('w:bidi')) is None:
        bidi = OxmlElement('w:bidi')
        grid = sectPr.find(qn('w:docGrid'))
        grid.addprevious(bidi) if grid is not None else sectPr.append(bidi)

def _fixed_layout(table):
    table.autofit = False
    tblPr = table._tbl.tblPr
    for el in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(el)
    lay = OxmlElement('w:tblLayout'); lay.set(qn('w:type'), 'fixed'); tblPr.append(lay)

def _shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'), 'clear'); shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def _repeat_header(row):
    row._tr.get_or_add_trPr().append(OxmlElement('w:tblHeader'))

def _widths(row, widths):
    for cell, w in zip(row.cells, widths):
        cell.width = w

def _font(run, name, size, bold):
    run.font.name = name; run.font.size = Pt(size); run.bold = bold
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts'); rPr.append(rFonts)
    rFonts.set(qn('w:cs'), name); rFonts.set(qn('w:ascii'), name); rFonts.set(qn('w:hAnsi'), name)

def he(doc, text, size=11, bold=False, space=6, center=False):
    p = doc.add_paragraph(); _set_rtl(p)          # bidi (RTL)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # justify aligns to the RTL start (right edge) for both headings and body;
        # an explicit jc=right flips to the left under a bidi section, so avoid it.
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); _font(r, HE_FONT, size, bold)
    r._element.get_or_add_rPr().append(OxmlElement('w:rtl'))
    p.paragraph_format.space_after = Pt(space)
    return p

def en(doc, text, size=11, bold=False, space=6, align=None):
    p = doc.add_paragraph()
    bidi = OxmlElement('w:bidi'); bidi.set(qn('w:val'), '0')  # force LTR even in an RTL section
    p._p.get_or_add_pPr().append(bidi)
    if align: p.alignment = align
    elif not (bold and size >= 13): p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); _font(r, EN_FONT, size, bold)
    p.paragraph_format.space_after = Pt(space)
    return p

# ---------------- shared course data ----------------
WEEKS_EN = [
    "Introduction to deep learning and framing a task as a network",
    "Tensors and data representation",
    "Multilayer perceptrons and backpropagation",
    "Data pipelines",
    "Loss functions and metrics",
    "Optimization (SGD, Adam)",
    "Regularization and generalization",
    "Convolutional networks I",
    "Convolutional networks II (normalization, residual connections)",
    "Recurrent networks (RNNs)",
    "LSTMs, GRUs and sequence tasks",
    "Representation learning (autoencoders, contrastive methods)",
    "Integration and transfer learning",
]
WEEKS_HE = [
    "מבוא ללמידה עמוקה וניסוח משימה כרשת",
    "טנזורים וייצוג נתונים",
    "פרספטרון רב-שכבתי והתפשטות לאחור",
    "צינורות נתונים",
    "פונקציות מחיר ומדדים",
    "אופטימיזציה (SGD, Adam)",
    "רגולריזציה והכללה",
    "רשתות קונבולוציה I",
    "רשתות קונבולוציה II (נרמול וחיבורים שאריתיים)",
    "רשתות נשנות (RNN)",
    "LSTM, GRU ומשימות רצף",
    "למידת ייצוגים (מקודדים אוטומטיים ושיטות ניגודיות)",
    "אינטגרציה ולמידת העברה",
]
BIB = [
    "Goodfellow, Ian, Yoshua Bengio, and Aaron Courville. Deep Learning. MIT Press, 2016.",
    "Prince, Simon J. D. Understanding Deep Learning. MIT Press, 2023.",
    "Zhang, Aston, Zachary C. Lipton, Mu Li, and Alexander J. Smola. Dive into Deep Learning. Cambridge University Press, 2023.",
    "Raschka, Sebastian, Yuxi Liu, and Vahid Mirjalili. Machine Learning with PyTorch and Scikit-Learn. Packt Publishing, 2022.",
    "Chollet, Francois. Deep Learning with Python. 2nd ed. Manning, 2021.",
]

def week_table(doc, header_left, header_right, weeks, rtl):
    NUM, TOPIC = (Cm(1.8), Cm(12.2)) if rtl else (Cm(2.9), Cm(11.8))   # match the examples
    t = doc.add_table(rows=1, cols=2); t.style = "Table Grid"
    _fixed_layout(t)
    hrow = t.rows[0]; _repeat_header(hrow)                 # repeat header across page breaks
    for c, txt in zip(hrow.cells, [header_left, header_right]):
        p = c.paragraphs[0]; r = p.add_run(txt); _font(r, HE_FONT if rtl else EN_FONT, 11, True)
        if rtl:
            _set_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r._element.get_or_add_rPr().append(OxmlElement('w:rtl'))
    _widths(hrow, [NUM, TOPIC])
    for i, wk in enumerate(weeks, 1):
        row = t.add_row()
        pn = row.cells[0].paragraphs[0]; pn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _font(pn.add_run(str(i)), EN_FONT, 11, False)
        ps = row.cells[1].paragraphs[0]; rs = ps.add_run(wk); _font(rs, HE_FONT if rtl else EN_FONT, 11, False)
        if rtl:
            _set_rtl(ps); ps.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            rs._element.get_or_add_rPr().append(OxmlElement('w:rtl'))
        _widths(row, [NUM, TOPIC])
    if rtl:
        t.alignment = 2
        _set_table_rtl(t)
    return t

# ---------------- 1. English syllabus ----------------
def build_syllabus_en():
    d = new_doc()
    en(d, "Introduction to Deep Learning", 20, False, 4, align=WD_ALIGN_PARAGRAPH.CENTER)
    en(d, "Lecture: 3 hours, practice: 2 hours", 11, True, 0)
    en(d, "5 hours, 4 credits", 11, False, 0)
    en(d, "Prerequisites: Machine Learning 63303, Probability 20021", 11, True, 10)
    en(d, "Course Objectives", 13, True, 4)
    en(d, "Deep learning underlies modern AI systems in vision, language, and beyond. This course gives a rigorous, hands-on foundation in neural networks: how to frame a machine-learning task in tensor terms, and how to build, train, and debug networks in PyTorch.")
    en(d, "It is a foundational course required of every computer-science student and is taught across all specializations. It provides the shared deep-learning base for continued specialization in artificial intelligence, including computer vision and large language models, and is valuable in its own right rather than only as a step toward later courses. By the end of the course, students can take a new task, frame it, build a suitable network, train it, diagnose failures, and improve it.")
    en(d, "Course content", 13, True, 4)
    en(d, "The course begins with the building blocks: tensors and data representation, the multilayer perceptron, and backpropagation with automatic differentiation. It then covers the training stack: data pipelines, loss functions and metrics, optimization (SGD, Adam), and regularization. The second half studies the core architecture families: convolutional networks (including normalization and residual connections), recurrent networks (RNNs, LSTMs, and GRUs), and representation learning with autoencoders and contrastive methods. The course concludes with transfer learning and an end-to-end workflow. It is highly practical, taught in PyTorch with weekly hands-on labs, and designed for working with an AI coding assistant.")
    en(d, "Student duties and grade components", 13, True, 4)
    en(d, "The course is project- and lab-based with no written exams. The final grade combines weekly labs (40%), a mid-term mini-project (20%), a final project with a short oral defense (35%), and participation (5%).")
    en(d, "Course of lessons", 13, True, 4)
    en(d, "Each week has three parts: a 3-hour lecture that develops the theory, a 2-hour practice lesson in which the instructor demonstrates implementations and works through examples, and a weekly lab set as homework. The labs follow a build, predict, and explain model: students may use an AI assistant for the build, but the graded work is predicting outcomes and explaining results.")
    en(d, "The order of the lessons (may change if required)", 11, True, 4)
    week_table(d, "Week", "Subject", WEEKS_EN, rtl=False)
    en(d, "", 6, False, 0)
    en(d, "Textbooks", 13, True, 4)
    for b in BIB: en(d, b, 11, False, 3)
    d.save(os.path.join(OUT, "syllabus_en.docx"))

# ---------------- 2. Hebrew syllabus ----------------
def build_syllabus_he():
    d = new_doc(); set_doc_rtl(d)
    he(d, "מבוא ללמידה עמוקה - Introduction to Deep Learning", 14, True, 4, center=True)
    he(d, "אופן הוראה: שיעור ותרגול.", 11, False, 0)
    he(d, "שעות שבועיות: הרצאה 3 שעות + תרגול 2 שעות, סה\"כ שעות – 5", 11, False, 0)
    he(d, "נקודות זכות: 4", 11, False, 0)
    he(d, "דרישות קדם: מבוא ללמידת מכונה 63303, הסתברות 20021", 11, False, 10)
    he(d, "מטרות הקורס", 13, True, 4)
    he(d, "כאשר מערכות בינה מלאכותית פועלות בתחומי הראייה הממוחשבת, השפה והדיבור, הן נשענות על רשתות נוירונים עמוקות. קורס זה מקנה בסיס מעמיק ומעשי ברשתות נוירונים: כיצד לנסח משימת למידת מכונה במונחים של טנזורים, וכיצד לבנות, לאמן ולנפות רשתות בעזרת PyTorch.")
    he(d, "זהו קורס יסוד הנדרש מכל סטודנט למדעי המחשב ונלמד בכל המסלולים. הוא מספק את הבסיס המשותף בלמידה עמוקה להמשך התמחות בבינה מלאכותית, ובכלל זה ראייה ממוחשבת ומודלי שפה גדולים, והוא בעל ערך בפני עצמו ולא רק כשלב לקראת קורסים מתקדמים. בסיום הקורס יוכלו הסטודנטים לקחת משימה חדשה, לנסח אותה, לבנות רשת מתאימה, לאמן אותה, לאבחן כשלים ולשפר אותה.")
    he(d, "תוכן הקורס", 13, True, 4)
    he(d, "הקורס נפתח באבני הבניין: טנזורים וייצוג נתונים, הפרספטרון הרב-שכבתי (MLP) ואלגוריתם ההתפשטות לאחור (backpropagation) עם גזירה אוטומטית. בהמשך נלמד את שלבי האימון: צינורות נתונים, פונקציות מחיר ומדדים, אופטימיזציה (SGD, Adam) ורגולריזציה. החצי השני עוסק במשפחות הארכיטקטורות המרכזיות: רשתות קונבולוציה (כולל נרמול וחיבורים שאריתיים), רשתות נשנות (RNN, LSTM ו-GRU), ולמידת ייצוגים בעזרת מקודדים אוטומטיים ושיטות ניגודיות (contrastive). הקורס מסתיים בלמידת העברה (transfer learning) ובתהליך עבודה מקצה לקצה. הקורס מעשי, נלמד ב-PyTorch עם תרגול שבועי, ומותאם לעבודה עם עוזר תכנות מבוסס בינה מלאכותית.")
    he(d, "חובות התלמידים ומרכיבי הציון", 13, True, 4)
    he(d, "הקורס מבוסס פרויקטים ומעבדות, ללא מבחן כתוב. הציון הסופי מורכב ממעבדות שבועיות (40%), פרויקט אמצע (20%), פרויקט סיום עם הגנה קצרה בעל-פה (35%), והשתתפות (5%).")
    he(d, "מהלך השיעורים", 13, True, 4)
    he(d, "כל שבוע כולל שלושה חלקים: הרצאה בת 3 שעות המפתחת את התיאוריה, שיעור תרגול בן שעתיים שבו המרצה מדגים יישומים ועובר על דוגמאות, ומעבדה הניתנת כשיעורי בית. המעבדות בנויות לפי מודל של בנייה, חיזוי והסבר: מותר להיעזר בעוזר בינה מלאכותית לשלב הבנייה, אך הציון ניתן על חיזוי התוצאות והסברן.")
    he(d, "שיטות ההוראה: הוראה פרונטלית מלווה במצגות ובהדגמות קוד.", 11, False, 2)
    he(d, "שימוש בטכנולוגיה: הדגמות ותרגול ב-Python ו-PyTorch, מחברות Colab, ושימוש בעוזר תכנות מבוסס בינה מלאכותית.", 11, False, 2)
    he(d, "מרצים אורחים: אין.", 11, False, 8)
    he(d, "תכנית הוראה מפורטת לכל השיעורים (סדר השיעורים צפוי להשתנות)", 11, True, 4)
    week_table(d, "שבוע", "נושאים", WEEKS_HE, rtl=True)
    he(d, "", 6, False, 0)
    he(d, "ביבליוגרפיה", 13, True, 4)
    for b in BIB: en(d, b, 11, False, 3, align=WD_ALIGN_PARAGRAPH.LEFT)
    d.save(os.path.join(OUT, "syllabus_he.docx"))

# ---------------- 3. Rationale ----------------
def build_rationale():
    d = new_doc(); set_doc_rtl(d)
    he(d, "מסמך רציונל לקורס מבוא ללמידה עמוקה", 14, True, 2, center=True)
    en(d, "Introduction to Deep Learning", 12, True, 10, align=WD_ALIGN_PARAGRAPH.CENTER)
    he(d, "הקורס עוסק ביסודות הלמידה העמוקה: ייצוג נתונים כטנזורים, בניית רשתות נוירונים ואימונן בעזרת PyTorch, והבנת הדינמיקה של אופטימיזציה, רגולריזציה ואבחון תהליך האימון. הקורס מקנה בסיס תיאורטי ומעשי כאחד.")
    he(d, "הקורס נדרש מכל הסטודנטים למדעי המחשב והוא קורס חובה בכל המסלולים. הוא מיועד לסטודנטים שסיימו קורס מבוא בלמידת מכונה, ומספק את הבסיס המשותף בלמידה עמוקה להמשך התמחות בבינה מלאכותית, ובכלל זה ראייה ממוחשבת ומודלי שפה גדולים. הקורס בעל ערך בפני עצמו ואינו רק שלב מקדים לקורסים מתקדמים.")
    he(d, "הנושאים שיילמדו בקורס: ניסוח משימות למידה כרשתות נוירונים; טנזורים וגזירה אוטומטית; רשתות MLP, רשתות קונבולוציה ורשתות נשנות (RNN, LSTM, GRU); פונקציות מחיר, אופטימיזציה ורגולריזציה; למידת ייצוגים ולמידת העברה; ותהליך עבודה מקצה לקצה ב-PyTorch.")
    he(d, "הקורס סוגר פער קיים בתכנית הלימודים: הוא מספק את הבסיס המשותף בלמידה עמוקה הנדרש לקורסי ההמשך, ומקנה עבודה מעשית עם כלי בינה מלאכותית מודרניים תוך שמירה על למידה אמיתית של החומר, כך שהסטודנט נדרש לחזות, להסביר ולהגן על עבודתו.")
    d.save(os.path.join(OUT, "rationale.docx"))

# ---------------- 4. Catalogue summary ----------------
def build_summary():
    d = new_doc(); set_doc_rtl(d)
    he(d, "תקצירים לידיעון", 14, True, 8, center=True)
    he(d, "מבוא ללמידה עמוקה", 13, True, 0, center=True)
    en(d, "Introduction to Deep Learning", 12, True, 8, align=WD_ALIGN_PARAGRAPH.CENTER)
    he(d, "אופן הוראה: שיעור ותרגול", 11, False, 0)
    he(d, "שעות שבועיות: הרצאה 3 שעות + תרגול 2 שעות, סה\"כ שעות – 5", 11, False, 0)
    he(d, "נקודות זכות: 4", 11, False, 0)
    he(d, "דרישות קדם: מבוא ללמידת מכונה 63303, הסתברות 20021", 11, False, 8)
    he(d, "קורס יסוד מעשי בלמידה עמוקה ב-PyTorch: ניסוח משימות כטנזורים, בניית רשתות נוירונים ואימונן, אופטימיזציה ורגולריזציה, רשתות קונבולוציה ונשנות, למידת ייצוגים ולמידת העברה. הקורס נדרש מכל הסטודנטים למדעי המחשב ומספק את הבסיס המשותף להמשך התמחות בבינה מלאכותית (ראייה ממוחשבת, מודלי שפה גדולים ועוד), ומשלב עבודה עם עוזר תכנות מבוסס בינה מלאכותית.")
    he(d, "נושאי הקורס: טנזורים והתפשטות לאחור; MLP, רשתות קונבולוציה ורשתות נשנות; פונקציות מחיר, אופטימיזציה ורגולריזציה; למידת ייצוגים ולמידת העברה.", 11, False, 12)
    en(d, "Introduction to Deep Learning", 13, True, 0)
    en(d, "Lecture and practice", 11, False, 0)
    en(d, "5 hours, 4 credits", 11, False, 0)
    en(d, "Prerequisites: Machine Learning 63303, Probability 20021", 11, False, 8)
    en(d, "A practical foundation course in deep learning with PyTorch: framing tasks as tensors, building and training neural networks, optimization and regularization, convolutional and recurrent networks, representation learning, and transfer learning. It is required of every computer-science student and provides the shared base for continued AI specialization (computer vision, large language models, and more); it integrates work with an AI coding assistant.")
    en(d, "Topics: tensors and backpropagation; MLPs, convolutional and recurrent networks; loss functions, optimization, and regularization; representation learning and transfer learning.")
    d.save(os.path.join(OUT, "catalogue_summary.docx"))

# ---------------- 5. Committee questionnaire ----------------
def build_questionnaire():
    d = new_doc(); set_doc_rtl(d)
    he(d, "שאלות למידע נוסף כהכנה לדיון בוועדת קוריקולום", 13, True, 6, center=True)
    he(d, "תאריך: _________", 11, False, 0)
    he(d, "פקולטה: מדעי המחשב            מחלקה: מדעי המחשב", 11, False, 8)
    rows = [
        ("שם הקורס בעברית: מבוא ללמידה עמוקה", "רציונל: קורס יסוד מעשי בלמידה עמוקה, חובה לכל סטודנטי מדעי המחשב, המספק בסיס משותף להמשך התמחות בבינה מלאכותית (ראייה ממוחשבת, מודלי שפה גדולים ועוד)."),
        ("קורס חובה או קורס בחירה (הקף בעיגול): חובה", "במסלול: כל מסלולי מדעי המחשב / לעיצוב: קורס עיוני"),
        ("נא לענות לגבי קורס חובה:", "הקורס אינו מחליף קורס קיים."),
        ("תחולה: משנה\"ל תשפ\"ז", ""),
        ("האם וכיצד ישפיע הקורס על הרכב הנ\"ז של תכנית הלימודים?", "קורס חובה (4 נ\"ז) בכל מסלולי מדעי המחשב."),
        ("קורסים דומים בפקולטה? (פרט)", "קיימים קורסי בחירה משיקים בלמידת מכונה; אין קורס יסוד ייעודי בלמידה עמוקה."),
        ("קורסים דומים במכון (פרט)", "אין קורס זהה במכון."),
        ("הערות נוספות:", "הקורס מעשי, מבוסס PyTorch, ומשלב עבודה עם עוזר תכנות מבוסס בינה מלאכותית תוך הערכה של חיזוי והסבר."),
    ]
    t = d.add_table(rows=0, cols=2); t.style = "Table Grid"; t.alignment = 2; _set_table_rtl(t); _fixed_layout(t)
    LABEL, CONTENT = Cm(6.5), Cm(9.1)
    for left, right in rows:
        row = t.add_row()
        for cell, txt in zip(row.cells, [left, right]):
            p = cell.paragraphs[0]; _set_rtl(p); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            r = p.add_run(txt); _font(r, HE_FONT, 11, False)
            r._element.get_or_add_rPr().append(OxmlElement('w:rtl'))
        _widths(row, [LABEL, CONTENT])
    d.save(os.path.join(OUT, "committee_questionnaire.docx"))

def main():
    os.makedirs(OUT, exist_ok=True)
    build_syllabus_en(); build_syllabus_he(); build_rationale(); build_summary(); build_questionnaire()
    print("HIT package written to hit-catalogue/:")
    for f in sorted(os.listdir(OUT)): print("  -", f)

if __name__ == "__main__":
    main()
